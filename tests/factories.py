"""Builders for the fixtures shared across the test suite.

Domain objects get sensible defaults, and parsers - being adapters
around third party libraries - are fed genuine PDF and DOCX bytes
instead of mocks.
"""

from __future__ import annotations

from collections.abc import Sequence
from datetime import UTC, datetime
from pathlib import Path

import docx

from devmind.domain.entities import ParsedDocument
from devmind.domain.value_objects import DocumentFormat, DocumentMetadata

CHECKSUM = "a1b2c3d4e5f60718" + "0" * 48
MODIFIED_AT = datetime(2026, 7, 29, 12, 0, tzinfo=UTC)


def make_metadata(**overrides: object) -> DocumentMetadata:
    """Build document metadata with sensible defaults for tests."""
    defaults: dict[str, object] = {
        "source_path": Path("data/documents/guide.md"),
        "document_format": DocumentFormat.MARKDOWN,
        "size_bytes": 128,
        "checksum": CHECKSUM,
        "modified_at": MODIFIED_AT,
    }
    return DocumentMetadata(**(defaults | overrides))  # type: ignore[arg-type]


def make_document(content: str, **overrides: object) -> ParsedDocument:
    """Build a parsed document carrying ``content``."""
    return ParsedDocument(content=content, metadata=make_metadata(**overrides))


_CATALOG_ID = 1
_PAGES_ID = 2
_FONT_ID = 3
_FIRST_PAGE_ID = 4


def _escape_pdf_text(text: str) -> str:
    """Escape the characters that delimit a PDF literal string."""
    return text.replace("\\", r"\\").replace("(", r"\(").replace(")", r"\)")


def build_pdf(
    pages: Sequence[str], *, title: str | None = None, author: str | None = None
) -> bytes:
    """Build a minimal, standards conformant PDF with a text layer.

    Args:
        pages: One text line per page.
        title: Optional title stored in the document information.
        author: Optional author stored in the document information.

    Returns:
        The complete PDF file as bytes.
    """
    if not pages:
        raise ValueError("A PDF needs at least one page.")

    page_ids = [_FIRST_PAGE_ID + index * 2 for index in range(len(pages))]
    content_ids = [page_id + 1 for page_id in page_ids]
    info_id = _FIRST_PAGE_ID + len(pages) * 2

    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    bodies: dict[int, bytes] = {
        _CATALOG_ID: f"<< /Type /Catalog /Pages {_PAGES_ID} 0 R >>".encode("ascii"),
        _PAGES_ID: f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>".encode("ascii"),
        _FONT_ID: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }

    for page_id, content_id, text in zip(page_ids, content_ids, pages, strict=True):
        bodies[page_id] = (
            f"<< /Type /Page /Parent {_PAGES_ID} 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 {_FONT_ID} 0 R >> >> "
            f"/Contents {content_id} 0 R >>"
        ).encode("ascii")
        stream = f"BT /F1 12 Tf 72 720 Td ({_escape_pdf_text(text)}) Tj ET".encode("ascii")
        bodies[content_id] = (
            f"<< /Length {len(stream)} >>\nstream\n".encode("ascii") + stream + b"\nendstream"
        )

    information = "".join(
        f"/{key} ({_escape_pdf_text(value)}) "
        for key, value in (("Title", title), ("Author", author))
        if value is not None
    )
    bodies[info_id] = f"<< {information}>>".encode("ascii")

    document = bytearray(b"%PDF-1.4\n")
    offsets: dict[int, int] = {}
    for object_id in sorted(bodies):
        offsets[object_id] = len(document)
        document += f"{object_id} 0 obj\n".encode("ascii")
        document += bodies[object_id]
        document += b"\nendobj\n"

    xref_offset = len(document)
    size = max(bodies) + 1
    document += f"xref\n0 {size}\n".encode("ascii")
    document += b"0000000000 65535 f \n"
    for object_id in range(1, size):
        document += f"{offsets[object_id]:010d} 00000 n \n".encode("ascii")
    document += (
        f"trailer\n<< /Size {size} /Root {_CATALOG_ID} 0 R /Info {info_id} 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    ).encode("ascii")
    return bytes(document)


def write_pdf(
    path: Path,
    pages: Sequence[str],
    *,
    title: str | None = None,
    author: str | None = None,
) -> Path:
    """Write a generated PDF to disk and return its path."""
    path.write_bytes(build_pdf(pages, title=title, author=author))
    return path


def write_docx(
    path: Path,
    paragraphs: Sequence[str] = (),
    *,
    table_rows: Sequence[Sequence[str]] = (),
    title: str | None = None,
    author: str | None = None,
    trailing_paragraphs: Sequence[str] = (),
) -> Path:
    """Write a DOCX file containing paragraphs, then a table, then more text.

    The optional trailing paragraphs make the reading order verifiable:
    a naive implementation that reads all paragraphs before all tables
    would place them before the table instead of after it.
    """
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row, values in zip(table.rows, table_rows, strict=True):
            for cell, value in zip(row.cells, values, strict=True):
                cell.text = value
    for text in trailing_paragraphs:
        document.add_paragraph(text)
    if title is not None:
        document.core_properties.title = title
    if author is not None:
        document.core_properties.author = author
    document.save(str(path))
    return path
